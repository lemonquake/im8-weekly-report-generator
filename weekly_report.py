#!/usr/bin/env python3
"""
Weekly Affiliate Report Generator (v2)
---------------------------------------
Reads a Social Snowball "Conversions" CSV export, aggregates the numbers, and
produces:
  1. A pretty terminal report (same as v1)
  2. A finished Word document (.docx) — either generated from scratch, or by
     filling {{PLACEHOLDERS}} in a template you provide.

USAGE
-----
Most common (auto-detect last week, auto-find newest CSV in folder, write a
fresh .docx report):

    python weekly_report.py

Pass a specific CSV:

    python weekly_report.py --csv conversions-export-1777527523.csv

By default the report uses the actual date span in the CSV (so a Mon-Fri
export gets a Mon-Fri report). Force a specific Mon-Sun week instead:

    python weekly_report.py --week 2026-04-20

Use a template with placeholders (see placeholders_reference.md):

    python weekly_report.py --template "WEEKLY REPORT STRUCTURE.docx"

Manually set the New Affiliates count from the SS Analytics dashboard:

    python weekly_report.py --new-affiliates 47

Exclude a specific top affiliate from the "without Top 1" math:

    python weekly_report.py --top1 "James DiNicolantonio"

Skip the .docx output entirely:

    python weekly_report.py --no-docx
"""

from __future__ import annotations

import argparse
import glob
import os
import re
import sys
from datetime import datetime, timedelta
from pathlib import Path

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── ANSI colors for terminal output ──────────────────────────────────────────
BOLD = "\033[1m"
CYAN = "\033[96m"
GREEN = "\033[92m"
YELLOW = "\033[93m"
RED = "\033[91m"
RESET = "\033[0m"
LINE = "-" * 60


def header(text):
    print(f"\n{BOLD}{CYAN}{text}{RESET}")
    print(LINE)


def row(label, value, indent=0):
    pad = "  " * indent
    print(f"{pad}{label:<35}{BOLD}{value}{RESET}")


# ── FORMAT HELPERS ───────────────────────────────────────────────────────────
def fmt_money(val: float) -> str:
    return f"${val:,.2f}"


def fmt_int(val) -> str:
    return f"{int(val):,}"


# ── LOAD & FILTER ────────────────────────────────────────────────────────────
def find_newest_csv(folder: Path) -> Path | None:
    """Find the newest conversions-export*.csv in the folder."""
    candidates = sorted(folder.glob("conversions-export*.csv"),
                        key=lambda p: p.stat().st_mtime, reverse=True)
    if candidates:
        return candidates[0]
    # Fallback: any csv
    candidates = sorted(folder.glob("*.csv"),
                        key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0] if candidates else None


def load_csv(csv_path: Path) -> pd.DataFrame:
    df = pd.read_csv(csv_path)
    df["Conversion Date"] = pd.to_datetime(df["Conversion Date"])
    df["Affiliate Name"] = (
        df["First Name"].astype(str).str.strip()
        + " "
        + df["Last Name"].astype(str).str.strip()
    )
    return df


def filter_week(df: pd.DataFrame, week_start: str | None) -> tuple[pd.DataFrame, datetime, datetime]:
    """Determine the reporting period.

    - If --week YYYY-MM-DD is given, use that Monday → Sunday (full 7-day window).
    - Otherwise, use the actual min/max Conversion Date in the CSV. This means
      a Mon-Fri export gets labeled as Mon-Fri, a Mon-Sun export as Mon-Sun.
    """
    if week_start:
        start = pd.Timestamp(week_start).normalize()
        end = start + timedelta(days=6, hours=23, minutes=59, seconds=59)
    else:
        start = df["Conversion Date"].min().normalize()
        # End of the day of the latest conversion
        latest = df["Conversion Date"].max()
        end = latest.normalize() + timedelta(hours=23, minutes=59, seconds=59)
    filtered = df[(df["Conversion Date"] >= start) & (df["Conversion Date"] <= end)].copy()
    return filtered, start, end


# ── AGGREGATE ────────────────────────────────────────────────────────────────
def aggregate(df: pd.DataFrame) -> pd.DataFrame:
    agg = (
        df.groupby("Affiliate Name", as_index=False)
        .agg(
            Revenue=("Revenue", "sum"),
            Commission=("Commission", "sum"),
            Conversions=("Order ID", "count"),
        )
    )
    agg["Profit"] = agg["Revenue"] - agg["Commission"]
    return agg.sort_values("Revenue", ascending=False).reset_index(drop=True)


def totals(agg: pd.DataFrame) -> dict:
    return {
        "revenue": float(agg["Revenue"].sum()),
        "commission": float(agg["Commission"].sum()),
        "conversions": int(agg["Conversions"].sum()),
        "profit": float(agg["Profit"].sum()),
        "active_affiliates": len(agg),
    }



# ── DOCX OUTPUT ──────────────────────────────────────────────────────────────
def _set_cell(cell, text, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(10)


def generate_docx_from_scratch(out_path: Path, ctx: dict, agg: pd.DataFrame,
                               top1_name: str):
    """Build a clean, professional weekly report .docx."""
    doc = Document()
    # Title
    title = doc.add_heading(f"Weekly Affiliate Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(ctx["week_label"])
    r.italic = True
    r.font.size = Pt(11)

    # ── Section: Overall totals ──────────────────────────────────────────────
    doc.add_heading("Overall Totals", level=1)
    t_all = ctx["totals_all"]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    rows_data = [
        ("Revenue", fmt_money(t_all["revenue"])),
        ("Commission", fmt_money(t_all["commission"])),
        ("Profit", fmt_money(t_all["profit"])),
        ("Conversions", fmt_int(t_all["conversions"])),
        ("Active Affiliates", fmt_int(t_all["active_affiliates"])),
    ]
    for label, val in rows_data:
        cells = table.add_row().cells
        _set_cell(cells[0], label, bold=True)
        _set_cell(cells[1], val, align="right")

    # ── Section: Without Top 1 ───────────────────────────────────────────────
    doc.add_heading(f"Without Top 1 ({top1_name})", level=1)
    t_ex = ctx["totals_excl"]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, val in [
        ("Revenue", fmt_money(t_ex["revenue"])),
        ("Commission", fmt_money(t_ex["commission"])),
        ("Profit", fmt_money(t_ex["profit"])),
        ("Conversions", fmt_int(t_ex["conversions"])),
        ("Active Affiliates", fmt_int(t_ex["active_affiliates"])),
    ]:
        cells = table.add_row().cells
        _set_cell(cells[0], label, bold=True)
        _set_cell(cells[1], val, align="right")

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  •  "
        f"Source: {ctx['csv_filename']}"
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(out_path)


def fill_template(template_path: Path, out_path: Path, ctx: dict,
                  agg: pd.DataFrame, top1_name: str):
    """Fill {{PLACEHOLDERS}} in a template docx with computed values.

    See placeholders_reference.md for the full list of supported placeholders.
    """
    doc = Document(template_path)
    t_all = ctx["totals_all"]
    t_ex = ctx["totals_excl"]

    mapping = {
        "{{WEEK_LABEL}}": ctx["week_label"],
        "{{WEEK_START}}": ctx["week_start"],
        "{{WEEK_END}}": ctx["week_end"],
        # Overall
        "{{REVENUE}}": fmt_money(t_all["revenue"]),
        "{{COMMISSION}}": fmt_money(t_all["commission"]),
        "{{PROFIT}}": fmt_money(t_all["profit"]),
        "{{CONVERSIONS}}": fmt_int(t_all["conversions"]),
        "{{ACTIVE_AFFILIATES}}": fmt_int(t_all["active_affiliates"]),
        # Excluding top 1
        "{{REVENUE_EXCL}}": fmt_money(t_ex["revenue"]),
        "{{COMMISSION_EXCL}}": fmt_money(t_ex["commission"]),
        "{{PROFIT_EXCL}}": fmt_money(t_ex["profit"]),
        "{{CONVERSIONS_EXCL}}": fmt_int(t_ex["conversions"]),
        "{{ACTIVE_AFFILIATES_EXCL}}": fmt_int(t_ex["active_affiliates"]),
        # Top 1 Name
        "{{TOP1_NAME}}": top1_name,
        # Date
        "{{GENERATED_AT}}": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }

    def replace_in_paragraph(p):
        text = p.text
        if not any(k in text for k in mapping):
            return
        for k, v in mapping.items():
            text = text.replace(k, v)
        # Replace by clearing runs and writing one run with the new text
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for r in table.rows:
            for cell in r.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    doc.save(out_path)


def generate_excel(out_path: Path, ctx: dict, top1_name: str):
    """Generate a clean Excel report with two sheets."""
    t_all = ctx["totals_all"]
    t_ex = ctx["totals_excl"]

    # ── Table 1: Overall Totals ──
    overall_data = {
        "Metric": ["Revenue", "Commission", "Profit", "Conversions", "Active Affiliates"],
        "Value": [
            t_all["revenue"],
            t_all["commission"],
            t_all["profit"],
            t_all["conversions"],
            t_all["active_affiliates"]
        ]
    }

    # ── Table 2: Without Top 1 ──
    excl_data = {
        "Metric": ["Revenue", "Commission", "Profit", "Conversions", "Active Affiliates"],
        "Value": [
            t_ex["revenue"],
            t_ex["commission"],
            t_ex["profit"],
            t_ex["conversions"],
            t_ex["active_affiliates"]
        ]
    }

    with pd.ExcelWriter(out_path, engine='openpyxl') as writer:
        df_overall = pd.DataFrame(overall_data)
        df_overall.to_excel(writer, sheet_name="Overall Totals", index=False)
        
        df_excl = pd.DataFrame(excl_data)
        df_excl.to_excel(writer, sheet_name=f"Without {top1_name[:20]}", index=False)


# ── TERMINAL OUTPUT ──────────────────────────────────────────────────────────
def print_terminal_report(ctx: dict, agg: pd.DataFrame, top1_name: str):
    print(f"\n{BOLD}{'='*60}{RESET}")
    print(f"{BOLD}  WEEKLY AFFILIATE REPORT  |  {ctx['week_label'].replace('–', '-')} {RESET}")
    print(f"{BOLD}{'='*60}{RESET}")

    t_all = ctx["totals_all"]
    header("OVERALL TOTALS")
    row("Revenue", fmt_money(t_all["revenue"]))
    row("Commission", fmt_money(t_all["commission"]))
    row("Profit", fmt_money(t_all["profit"]))
    row("Conversions", fmt_int(t_all["conversions"]))

    t_ex = ctx["totals_excl"]
    header(f"WITHOUT TOP 1 ({top1_name})")
    row("Revenue", fmt_money(t_ex["revenue"]))
    row("Commission", fmt_money(t_ex["commission"]))
    row("Profit", fmt_money(t_ex["profit"]))
    row("Conversions", fmt_int(t_ex["conversions"]))

    header("CROSS-CHECK AGAINST SOCIAL SNOWBALL ANALYTICS")
    print("  These should match the Analytics tiles for the same week:")
    print(f"    Conversions  -> {fmt_int(t_all['conversions'])}")
    print(f"    Commissions  -> {fmt_money(t_all['commission'])}")
    print(f"    Revenue      -> {fmt_money(t_all['revenue'])}")
    print(f"\n{BOLD}{'='*60}{RESET}\n")


# ── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="Generate the weekly affiliate report (terminal + .docx).")
    parser.add_argument("csv_pos", nargs="?",
                        help="Path to CSV (positional, optional)")
    parser.add_argument("--csv", help="Path to the Social Snowball CSV export")
    parser.add_argument("--top1",
                        help='Full name of top affiliate (default: auto-detect by Revenue)')
    parser.add_argument("--template",
                        help="Path to a .docx template with {{PLACEHOLDERS}} to fill in")
    parser.add_argument("--out", help="Path for the generated .docx (default: auto)")
    parser.add_argument("--no-docx", action="store_true",
                        help="Skip Word document output")
    args = parser.parse_args()

    here = Path(__file__).resolve().parent

    # ── Resolve CSV ──────────────────────────────────────────────────────────
    csv_arg = args.csv or args.csv_pos
    if csv_arg:
        csv_path = Path(csv_arg)
        if not csv_path.is_absolute():
            csv_path = here / csv_path
    else:
        csv_path = find_newest_csv(here)
        if csv_path is None:
            print(f"{RED}No CSV found in {here}. Drop a conversions-export*.csv "
                  f"in this folder or pass --csv.{RESET}")
            sys.exit(1)
        print(f"{CYAN}Auto-selected CSV: {csv_path.name}{RESET}")

    if not csv_path.exists():
        print(f"{RED}CSV not found: {csv_path}{RESET}")
        sys.exit(1)

    # ── Load + filter ────────────────────────────────────────────────────────
    df_all = load_csv(csv_path)
    week_df, start, end = filter_week(df_all, None)
    if week_df.empty:
        print(f"\n{YELLOW}No data found for week {start.date()} – {end.date()}.{RESET}")
        print(f"  CSV spans: {df_all['Conversion Date'].min().date()} "
              f"to {df_all['Conversion Date'].max().date()}")
        sys.exit(1)

    agg = aggregate(week_df)

    # ── Top 1 ────────────────────────────────────────────────────────────────
    if args.top1:
        top1_name = args.top1.strip()
        if top1_name not in agg["Affiliate Name"].values:
            print(f"\n{YELLOW}'{top1_name}' not found in this week. "
                  f"Available: {', '.join(agg['Affiliate Name'].tolist())}{RESET}")
            sys.exit(1)
    else:
        top1_name = agg.iloc[0]["Affiliate Name"]

    # ── Build context ────────────────────────────────────────────────────────
    week_label = f"{start.strftime('%b %d')} – {end.strftime('%b %d, %Y')}"
    ctx = {
        "week_label": week_label,
        "week_start": start.strftime("%Y-%m-%d"),
        "week_end": end.strftime("%Y-%m-%d"),
        "totals_all": totals(agg),
        "totals_excl": totals(agg[agg["Affiliate Name"] != top1_name]),
        "csv_filename": csv_path.name,
    }

    # ── Print terminal report ────────────────────────────────────────────────
    print_terminal_report(ctx, agg, top1_name)

    # ── Write docx ───────────────────────────────────────────────────────────
    if args.no_docx:
        return

    if not DOCX_AVAILABLE:
        print(f"{YELLOW}python-docx not installed; skipping .docx output. "
              f"Run:  pip install python-docx{RESET}")
        return

    out_path = Path(args.out) if args.out else (
        here / f"Weekly_Report_{start.strftime('%Y-%m-%d')}.docx"
    )

    if args.template:
        tpl = Path(args.template)
        if not tpl.is_absolute():
            tpl = here / tpl
        if not tpl.exists():
            print(f"{RED}Template not found: {tpl}{RESET}")
            sys.exit(1)
        fill_template(tpl, out_path, ctx, agg, top1_name)
        print(f"{GREEN}Wrote (from template): {out_path}{RESET}")
    else:
        generate_docx_from_scratch(out_path, ctx, agg, top1_name)
        print(f"{GREEN}Wrote: {out_path}{RESET}")


if __name__ == "__main__":
    main()
